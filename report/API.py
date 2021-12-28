from docx import Document
from docx.shared import Inches
from .Function import Function
from .Attribut import Attribut


class API:
    coder = ""
    desc = ""
    functions = []
    def __init__(self,):
        self.functions = []
    def get_report(self):
        document = Document()
        document.add_heading('API documentation', 0)
        document.add_heading('This API is maded by :', level=1)
        coder = self.coder.replace("," , " and ")
        p = document.add_paragraph(coder)
        document.add_heading('Description :', level=1)
        p = document.add_paragraph(self.desc)

        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'URL'
        hdr_cells[1].text = 'INPUT'
        hdr_cells[2].text = 'OUTPUT'
        hdr_cells[3].text = 'Definition'
        table.style ="LightGrid-Accent1"
        for f in self.functions:
            cells = table.add_row().cells
            cells[0].text = f.url
            cells[1].text = f.get_att()
            cells[2].text = f.returnn
            cells[3].text = f.desc
        document.save('report.docx')
    
    @staticmethod
    def create(name): 
        lines = API.get_lines(name)
        api = API()
        for line in lines:
            if "madeby::" in line : 
                api.coder = line.replace("#", "").replace("madeby::", "").replace("\n", "")
            elif("desc::" in line):
                api.desc = line.replace("desc::", "")
            elif ("@app.route" in line) :
                typee = ""
                if "method" not in line:
                    typee = "GET, POST, DELETE, PUT"
                else: 
                    typee = line.split("[")[1].replace("]" ,"").replace(")", "")        
                url = line.split("'")[1]
                fun = Function(typee, url)
                api.functions.append(fun)
            elif (":fn:" in line):
                l = line.split("\n")
                fun.desc = l[0].replace(":fn:","")
                for k in l:
                    if (":param" in k): 
                        c = k.split(":")
                        p = Attribut(c[1].split(" ")[1], c[2])
                        fun.attributs.append(p)
                    if ":return:" in k : 
                        fun.returnn = k.replace(":return:", "")
        return api

    @staticmethod
    def get_lines(name): 
        file = open(name, "r")
        lines = []
        comment = ""
        begin_comment = False
        for line in file:
                if "'''" in line :
                    if not(begin_comment):
                        begin_comment = True
                        continue
                    else:
                        line.replace("'''", '')
                        begin_comment = False
                        lines.append(comment)
                        comment = ""
                else:
                    if begin_comment : 
                        comment += line + " "
                    else:
                        lines.append(line)
        return lines
